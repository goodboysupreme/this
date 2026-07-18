"""Resume screener endpoints (Pillar B).

Flow: parse upload in memory → try DeepSeek (only when a key is configured)
→ fall back to the deterministic offline scorer on any failure. Uploaded
bytes are never written to disk and resume text is never stored in the DB.
"""

from __future__ import annotations

import re
import uuid
from pathlib import Path

import docx
from fastapi import APIRouter, File, Form, HTTPException, UploadFile
from fastapi.responses import FileResponse

from ..config import settings
from ..schemas import JdMatchResult, ResumeAnalysis
from ..services import ats_fallback, deepseek_client, resume_parser

router = APIRouter(prefix="/resume", tags=["resume"])

MAX_UPLOAD_BYTES = 5 * 1024 * 1024
EXPORT_DIR = Path(__file__).resolve().parents[2] / "uploads" / "exports"
_DOWNLOAD_ID_RE = re.compile(r"^[0-9a-f]{32}$")


def _parse_upload(file: UploadFile, data: bytes) -> str:
    """Validate + parse an uploaded resume; return plain text."""
    if not data:
        raise HTTPException(status_code=400, detail="empty file uploaded")
    if len(data) > MAX_UPLOAD_BYTES:
        raise HTTPException(status_code=413, detail="file exceeds the 5MB limit")
    try:
        parsed = resume_parser.parse_resume(file.filename or "resume", data)
    except ValueError as exc:
        raise HTTPException(status_code=400, detail=str(exc)) from exc
    except Exception as exc:
        raise HTTPException(
            status_code=422, detail="could not extract text from the file"
        ) from exc
    if not parsed["text"].strip():
        raise HTTPException(status_code=422, detail="no text found in the file")
    return parsed["text"]


def _clamp_int(value, default: int = 0) -> int:
    try:
        return max(0, min(100, int(round(float(value)))))
    except (TypeError, ValueError):
        return default


def _str_list(value) -> list[str]:
    if not isinstance(value, (list, tuple)):
        return []
    return sorted({str(v) for v in value if isinstance(v, (str, int, float))})


def _normalize_deepseek_analysis(raw: dict) -> dict:
    """Normalize DeepSeek's JSON into the unified analysis shape."""
    if not isinstance(raw, dict):
        raise ValueError("DeepSeek analysis payload is not a JSON object")
    sections = []
    for item in raw.get("sections") or []:
        if not isinstance(item, dict) or "name" not in item:
            continue
        sections.append(
            {
                "name": str(item["name"]),
                "score": _clamp_int(item.get("score")),
                "issues": [str(i) for i in (item.get("issues") or [])],
                "suggestions": [str(s) for s in (item.get("suggestions") or [])],
            }
        )
    checks = []
    for item in raw.get("ats_checks") or []:
        if not isinstance(item, dict) or "name" not in item:
            continue
        checks.append(
            {
                "name": str(item["name"]),
                "passed": bool(item.get("passed")),
                "detail": str(item.get("detail", "")),
            }
        )
    return {
        "engine": "deepseek",
        "overall_score": _clamp_int(raw.get("overall_score")),
        "sections": sections,
        "skills_detected": _str_list(raw.get("skills_detected")),
        "ats_checks": checks,
        "summary": str(raw.get("summary", "")),
    }


def _normalize_deepseek_jd(raw: dict) -> dict:
    """Normalize DeepSeek's JSON into the unified JD-match shape."""
    if not isinstance(raw, dict):
        raise ValueError("DeepSeek JD payload is not a JSON object")
    bullets = [
        {"original": str(b.get("original", "")), "suggested": str(b.get("suggested", ""))}
        for b in (raw.get("tailored_bullets") or [])
        if isinstance(b, dict)
    ]
    projects = []
    for p in raw.get("suggested_projects") or []:
        if not isinstance(p, dict):
            continue
        projects.append(
            {
                "title": str(p.get("title", "")),
                "stack": [str(s) for s in (p.get("stack") or [])],
                "why": str(p.get("why", "")),
                "talking_points": [str(t) for t in (p.get("talking_points") or [])],
            }
        )
    return {
        "match_score": _clamp_int(raw.get("match_score")),
        "matched_keywords": _str_list(raw.get("matched_keywords")),
        "missing_keywords": _str_list(raw.get("missing_keywords")),
        "tailored_bullets": bullets,
        "suggested_projects": projects,
    }


def _write_tailored_docx(download_id: str, text: str, result: dict) -> None:
    """Render a clean single-column ATS-friendly resume docx."""
    EXPORT_DIR.mkdir(parents=True, exist_ok=True)
    lines = [line.strip() for line in text.splitlines() if line.strip()]
    name = lines[0] if lines else "Candidate"

    document = docx.Document()
    document.add_heading(name, level=0)

    contact_bits = []
    for regex in (ats_fallback.EMAIL_RE, ats_fallback.PHONE_RE):
        found = regex.search(text)
        if found:
            contact_bits.append(found.group(0).strip())
    for link in re.findall(r"(?i)\S*(?:linkedin\.com|github\.com)\S*", text):
        contact_bits.append(link.strip(".,)"))
    if contact_bits:
        document.add_paragraph(" | ".join(dict.fromkeys(contact_bits)))

    skills = ats_fallback.extract_skills(text)
    if skills:
        document.add_heading("Skills", level=1)
        document.add_paragraph(", ".join(skills))

    bullets = ats_fallback.extract_bullets(text)
    tailored_by_original = {
        b["original"]: b["suggested"] for b in result.get("tailored_bullets", [])
    }
    if bullets:
        document.add_heading("Experience & Projects", level=1)
        for bullet in bullets:
            document.add_paragraph(
                tailored_by_original.get(bullet, bullet), style="List Bullet"
            )

    projects = result.get("suggested_projects", [])
    if projects:
        document.add_heading("Suggested Projects (JD-targeted)", level=1)
        for proj in projects:
            document.add_paragraph(proj["title"], style="List Bullet")
            if proj.get("stack"):
                document.add_paragraph(f"Stack: {', '.join(proj['stack'])}")

    document.save(EXPORT_DIR / f"{download_id}.docx")


@router.post("/analyze", response_model=ResumeAnalysis)
def analyze_resume(file: UploadFile = File(...)) -> dict:
    text = _parse_upload(file, file.file.read())
    if settings.deepseek_api_key:
        try:
            return _normalize_deepseek_analysis(deepseek_client.analyze_resume(text))
        except Exception:
            pass  # any failure → offline engine
    result = ats_fallback.analyze(text)
    result["engine"] = "fallback"
    return result


@router.post("/jd-match", response_model=JdMatchResult)
def jd_match_resume(file: UploadFile = File(...), jd_text: str = Form(...)) -> dict:
    text = _parse_upload(file, file.file.read())
    if not jd_text.strip():
        raise HTTPException(status_code=400, detail="jd_text must not be empty")

    result = None
    engine = "fallback"
    if settings.deepseek_api_key:
        try:
            result = _normalize_deepseek_jd(deepseek_client.jd_match(text, jd_text))
            engine = "deepseek"
        except Exception:
            result = None
    if result is None:
        result = ats_fallback.jd_match(text, jd_text)

    download_id = uuid.uuid4().hex
    _write_tailored_docx(download_id, text, result)
    result["engine"] = engine
    result["download_id"] = download_id
    return result


@router.get("/export/{download_id}")
def export_tailored_resume(download_id: str) -> FileResponse:
    if not _DOWNLOAD_ID_RE.fullmatch(download_id):
        raise HTTPException(status_code=404, detail="export not found")
    path = EXPORT_DIR / f"{download_id}.docx"
    if not path.is_file():
        raise HTTPException(status_code=404, detail="export not found")
    return FileResponse(
        path,
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        filename="tailored_resume.docx",
    )
